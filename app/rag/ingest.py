from pathlib import Path

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument

from app.config import settings


def load_docx(path: Path) -> list[Document]:
  doc = DocxDocument(str(path))
  text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
  return [Document(page_content=text, metadata={"source": path.name})]


def load_documents(file_path: Path) -> list[Document]:
  suffix = file_path.suffix.lower()

  if suffix == ".pdf":
    loader = PyPDFLoader(str(file_path))
    docs = loader.load()
    for doc in docs:
      doc.metadata["source"] = file_path.name
    return docs

  if suffix == ".txt":
    loader = TextLoader(str(file_path), encoding="utf-8")
    docs = loader.load()
    for doc in docs:
      doc.metadata["source"] = file_path.name
    return docs

  if suffix == ".docx":
    return load_docx(file_path)

  raise ValueError(f"Unsupported file type: {suffix}. Use .pdf, .txt, or .docx")


def split_documents(documents: list[Document]) -> list[Document]:
  splitter = RecursiveCharacterTextSplitter(
    chunk_size=settings.chunk_size,
    chunk_overlap=settings.chunk_overlap,
  )
  return splitter.split_documents(documents)
